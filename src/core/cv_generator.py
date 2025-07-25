"""
CV generator for the smart CV writer service
"""
import os
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..models.cv_data import CVData, CVTemplate

logger = logging.getLogger(__name__)


class CVGenerator:
    """Generates CV documents in various formats"""
    
    def __init__(self, template_style: str = "modern"):
        self.template_style = template_style
        self.template_config = self._get_template_config(template_style)
    
    def generate_pdf(self, cv_data: CVData, output_path: str) -> str:
        """Generate PDF CV"""
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=self.template_config["margin"] * inch,
                leftMargin=self.template_config["margin"] * inch,
                topMargin=self.template_config["margin"] * inch,
                bottomMargin=self.template_config["margin"] * inch
            )
            
            # Build story (content)
            story = []
            
            # Add header
            story.extend(self._create_header(cv_data))
            
            # Add summary
            if cv_data.summary:
                story.extend(self._create_summary_section(cv_data.summary))
            
            # Add experience
            if cv_data.employment_history:
                story.extend(self._create_experience_section(cv_data.employment_history))
            
            # Add skills
            if cv_data.technical_skills or cv_data.soft_skills:
                story.extend(self._create_skills_section(cv_data))
            
            # Add education
            if cv_data.education:
                story.extend(self._create_education_section(cv_data.education))
            
            # Add projects
            if cv_data.projects:
                story.extend(self._create_projects_section(cv_data.projects))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF CV generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF CV: {str(e)}")
            raise
    
    def generate_docx(self, cv_data: CVData, output_path: str) -> str:
        """Generate DOCX CV"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed")
        
        try:
            doc = Document()
            
            # Set up document styles
            self._setup_docx_styles(doc)
            
            # Add header
            self._add_docx_header(doc, cv_data)
            
            # Add summary
            if cv_data.summary:
                self._add_docx_summary(doc, cv_data.summary)
            
            # Add experience
            if cv_data.employment_history:
                self._add_docx_experience(doc, cv_data.employment_history)
            
            # Add skills
            if cv_data.technical_skills or cv_data.soft_skills:
                self._add_docx_skills(doc, cv_data)
            
            # Add education
            if cv_data.education:
                self._add_docx_education(doc, cv_data.education)
            
            # Add projects
            if cv_data.projects:
                self._add_docx_projects(doc, cv_data.projects)
            
            # Save document
            doc.save(output_path)
            
            logger.info(f"DOCX CV generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX CV: {str(e)}")
            raise
    
    def generate_text(self, cv_data: CVData, output_path: str) -> str:
        """Generate plain text CV"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Header
                f.write(self._create_text_header(cv_data))
                f.write('\n\n')
                
                # Summary
                if cv_data.summary:
                    f.write("PROFESSIONAL SUMMARY\n")
                    f.write("=" * 50 + '\n')
                    f.write(cv_data.summary)
                    f.write('\n\n')
                
                # Experience
                if cv_data.employment_history:
                    f.write("PROFESSIONAL EXPERIENCE\n")
                    f.write("=" * 50 + '\n')
                    for job in cv_data.employment_history:
                        f.write(f"{job.position} at {job.company}\n")
                        f.write(f"{job.start_date} - {job.end_date}\n")
                        f.write(f"Location: {job.location}\n\n")
                        
                        if job.description:
                            f.write("Responsibilities:\n")
                            for desc in job.description:
                                f.write(f"• {desc}\n")
                            f.write('\n')
                        
                        if job.achievements:
                            f.write("Key Achievements:\n")
                            for achievement in job.achievements:
                                f.write(f"• {achievement}\n")
                            f.write('\n')
                
                # Skills
                if cv_data.technical_skills or cv_data.soft_skills:
                    f.write("SKILLS\n")
                    f.write("=" * 50 + '\n')
                    
                    if cv_data.technical_skills:
                        f.write("Technical Skills:\n")
                        for skill in cv_data.technical_skills:
                            f.write(f"• {skill.name}")
                            if skill.proficiency:
                                f.write(f" ({skill.proficiency})")
                            f.write('\n')
                        f.write('\n')
                    
                    if cv_data.soft_skills:
                        f.write("Soft Skills:\n")
                        for skill in cv_data.soft_skills:
                            f.write(f"• {skill.name}")
                            if skill.proficiency:
                                f.write(f" ({skill.proficiency})")
                            f.write('\n')
                        f.write('\n')
                
                # Education
                if cv_data.education:
                    f.write("EDUCATION\n")
                    f.write("=" * 50 + '\n')
                    for edu in cv_data.education:
                        f.write(f"{edu.degree} in {edu.field_of_study}\n")
                        f.write(f"{edu.institution}\n")
                        f.write(f"Graduated: {edu.graduation_date}\n")
                        if edu.gpa:
                            f.write(f"GPA: {edu.gpa}\n")
                        f.write('\n')
                
                # Projects
                if cv_data.projects:
                    f.write("PROJECTS\n")
                    f.write("=" * 50 + '\n')
                    for project in cv_data.projects:
                        f.write(f"{project.name}\n")
                        f.write(f"Technologies: {', '.join(project.technologies)}\n")
                        f.write(f"Description: {project.description}\n")
                        if project.url:
                            f.write(f"URL: {project.url}\n")
                        f.write('\n')
            
            logger.info(f"Text CV generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating text CV: {str(e)}")
            raise
    
    def _get_template_config(self, style: str) -> Dict[str, Any]:
        """Get template configuration"""
        templates = {
            "modern": {
                "font_family": "Helvetica",
                "header_color": colors.HexColor("#2c3e50"),
                "section_color": colors.HexColor("#34495e"),
                "accent_color": colors.HexColor("#3498db"),
                "line_spacing": 1.15,
                "margin": 0.75
            },
            "professional": {
                "font_family": "Times-Roman",
                "header_color": colors.black,
                "section_color": colors.HexColor("#333333"),
                "accent_color": colors.HexColor("#666666"),
                "line_spacing": 1.2,
                "margin": 1.0
            },
            "creative": {
                "font_family": "Helvetica",
                "header_color": colors.HexColor("#1a237e"),
                "section_color": colors.HexColor("#303f9f"),
                "accent_color": colors.HexColor("#3f51b5"),
                "line_spacing": 1.1,
                "margin": 0.5
            }
        }
        
        return templates.get(style, templates["modern"])
    
    def _create_header(self, cv_data: CVData) -> List:
        """Create PDF header section"""
        story = []
        styles = getSampleStyleSheet()
        
        # Name
        name_style = ParagraphStyle(
            'Name',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=self.template_config["header_color"],
            alignment=TA_CENTER,
            spaceAfter=12
        )
        story.append(Paragraph(cv_data.contact_info.full_name, name_style))
        
        # Contact info
        contact_info = []
        if cv_data.contact_info.email:
            contact_info.append(cv_data.contact_info.email)
        if cv_data.contact_info.phone:
            contact_info.append(cv_data.contact_info.phone)
        if cv_data.contact_info.location:
            contact_info.append(cv_data.contact_info.location)
        
        if contact_info:
            contact_style = ParagraphStyle(
                'Contact',
                parent=styles['Normal'],
                fontSize=10,
                textColor=self.template_config["section_color"],
                alignment=TA_CENTER,
                spaceAfter=12
            )
            story.append(Paragraph(' | '.join(contact_info), contact_style))
        
        # Online presence
        online_info = []
        if cv_data.contact_info.linkedin:
            online_info.append(f"LinkedIn: {cv_data.contact_info.linkedin}")
        if cv_data.contact_info.github:
            online_info.append(f"GitHub: {cv_data.contact_info.github}")
        if cv_data.contact_info.portfolio:
            online_info.append(f"Portfolio: {cv_data.contact_info.portfolio}")
        
        if online_info:
            online_style = ParagraphStyle(
                'Online',
                parent=styles['Normal'],
                fontSize=9,
                textColor=self.template_config["accent_color"],
                alignment=TA_CENTER,
                spaceAfter=20
            )
            story.append(Paragraph(' | '.join(online_info), online_style))
        
        return story
    
    def _create_summary_section(self, summary: str) -> List:
        """Create PDF summary section"""
        story = []
        styles = getSampleStyleSheet()
        
        # Section header
        header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=self.template_config["section_color"],
            spaceAfter=6,
            spaceBefore=12
        )
        story.append(Paragraph("PROFESSIONAL SUMMARY", header_style))
        
        # Summary content
        summary_style = ParagraphStyle(
            'Summary',
            parent=styles['Normal'],
            fontSize=10,
            leading=12,
            spaceAfter=12
        )
        story.append(Paragraph(summary, summary_style))
        
        return story
    
    def _create_experience_section(self, employment_history: List) -> List:
        """Create PDF experience section"""
        story = []
        styles = getSampleStyleSheet()
        
        # Section header
        header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=self.template_config["section_color"],
            spaceAfter=6,
            spaceBefore=12
        )
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", header_style))
        
        # Job entries
        for job in employment_history:
            # Job title and company
            title_style = ParagraphStyle(
                'JobTitle',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=self.template_config["header_color"],
                spaceAfter=2
            )
            story.append(Paragraph(f"{job.position} at {job.company}", title_style))
            
            # Dates and location
            date_style = ParagraphStyle(
                'JobDates',
                parent=styles['Normal'],
                fontSize=9,
                textColor=self.template_config["accent_color"],
                spaceAfter=6
            )
            date_text = f"{job.start_date} - {job.end_date}"
            if job.location:
                date_text += f" | {job.location}"
            story.append(Paragraph(date_text, date_style))
            
            # Description
            if job.description:
                for desc in job.description:
                    desc_style = ParagraphStyle(
                        'JobDescription',
                        parent=styles['Normal'],
                        fontSize=10,
                        leftIndent=20,
                        spaceAfter=2
                    )
                    story.append(Paragraph(f"• {desc}", desc_style))
            
            # Achievements
            if job.achievements:
                for achievement in job.achievements:
                    achievement_style = ParagraphStyle(
                        'JobAchievement',
                        parent=styles['Normal'],
                        fontSize=10,
                        leftIndent=20,
                        spaceAfter=2
                    )
                    story.append(Paragraph(f"• {achievement}", achievement_style))
            
            story.append(Spacer(1, 12))
        
        return story
    
    def _create_skills_section(self, cv_data: CVData) -> List:
        """Create PDF skills section"""
        story = []
        styles = getSampleStyleSheet()
        
        # Section header
        header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=self.template_config["section_color"],
            spaceAfter=6,
            spaceBefore=12
        )
        story.append(Paragraph("SKILLS", header_style))
        
        # Technical skills
        if cv_data.technical_skills:
            tech_header_style = ParagraphStyle(
                'SkillHeader',
                parent=styles['Heading3'],
                fontSize=11,
                textColor=self.template_config["header_color"],
                spaceAfter=4
            )
            story.append(Paragraph("Technical Skills:", tech_header_style))
            
            tech_skills = []
            for skill in cv_data.technical_skills:
                skill_text = skill.name
                if skill.proficiency:
                    skill_text += f" ({skill.proficiency})"
                tech_skills.append(skill_text)
            
            skill_style = ParagraphStyle(
                'Skill',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=2
            )
            story.append(Paragraph(', '.join(tech_skills), skill_style))
            story.append(Spacer(1, 8))
        
        # Soft skills
        if cv_data.soft_skills:
            soft_header_style = ParagraphStyle(
                'SkillHeader',
                parent=styles['Heading3'],
                fontSize=11,
                textColor=self.template_config["header_color"],
                spaceAfter=4
            )
            story.append(Paragraph("Soft Skills:", soft_header_style))
            
            soft_skills = []
            for skill in cv_data.soft_skills:
                skill_text = skill.name
                if skill.proficiency:
                    skill_text += f" ({skill.proficiency})"
                soft_skills.append(skill_text)
            
            skill_style = ParagraphStyle(
                'Skill',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=2
            )
            story.append(Paragraph(', '.join(soft_skills), skill_style))
        
        return story
    
    def _create_education_section(self, education: List) -> List:
        """Create PDF education section"""
        story = []
        styles = getSampleStyleSheet()
        
        # Section header
        header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=self.template_config["section_color"],
            spaceAfter=6,
            spaceBefore=12
        )
        story.append(Paragraph("EDUCATION", header_style))
        
        # Education entries
        for edu in education:
            # Degree and institution
            degree_style = ParagraphStyle(
                'Degree',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=self.template_config["header_color"],
                spaceAfter=2
            )
            story.append(Paragraph(f"{edu.degree} in {edu.field_of_study}", degree_style))
            
            # Institution and dates
            institution_style = ParagraphStyle(
                'Institution',
                parent=styles['Normal'],
                fontSize=10,
                textColor=self.template_config["accent_color"],
                spaceAfter=2
            )
            story.append(Paragraph(f"{edu.institution} | {edu.graduation_date}", institution_style))
            
            if edu.gpa:
                gpa_style = ParagraphStyle(
                    'GPA',
                    parent=styles['Normal'],
                    fontSize=9,
                    spaceAfter=8
                )
                story.append(Paragraph(f"GPA: {edu.gpa}", gpa_style))
            
            story.append(Spacer(1, 8))
        
        return story
    
    def _create_projects_section(self, projects: List) -> List:
        """Create PDF projects section"""
        story = []
        styles = getSampleStyleSheet()
        
        # Section header
        header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=self.template_config["section_color"],
            spaceAfter=6,
            spaceBefore=12
        )
        story.append(Paragraph("PROJECTS", header_style))
        
        # Project entries
        for project in projects:
            # Project name
            name_style = ParagraphStyle(
                'ProjectName',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=self.template_config["header_color"],
                spaceAfter=2
            )
            story.append(Paragraph(project.name, name_style))
            
            # Technologies
            if project.technologies:
                tech_style = ParagraphStyle(
                    'ProjectTech',
                    parent=styles['Normal'],
                    fontSize=9,
                    textColor=self.template_config["accent_color"],
                    spaceAfter=4
                )
                story.append(Paragraph(f"Technologies: {', '.join(project.technologies)}", tech_style))
            
            # Description
            if project.description:
                desc_style = ParagraphStyle(
                    'ProjectDescription',
                    parent=styles['Normal'],
                    fontSize=10,
                    spaceAfter=8
                )
                story.append(Paragraph(project.description, desc_style))
        
        return story
    
    def _setup_docx_styles(self, doc: Document):
        """Set up DOCX document styles"""
        # This would configure document styles
        pass
    
    def _add_docx_header(self, doc: Document, cv_data: CVData):
        """Add header to DOCX document"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(cv_data.contact_info.full_name)
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_info = []
        if cv_data.contact_info.email:
            contact_info.append(cv_data.contact_info.email)
        if cv_data.contact_info.phone:
            contact_info.append(cv_data.contact_info.phone)
        if cv_data.contact_info.location:
            contact_info.append(cv_data.contact_info.location)
        
        if contact_info:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(' | '.join(contact_info))
            contact_run.font.size = Pt(10)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
    
    def _add_docx_summary(self, doc: Document, summary: str):
        """Add summary to DOCX document"""
        doc.add_heading("PROFESSIONAL SUMMARY", level=1)
        doc.add_paragraph(summary)
        doc.add_paragraph()  # Spacing
    
    def _add_docx_experience(self, doc: Document, employment_history: List):
        """Add experience to DOCX document"""
        doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
        
        for job in employment_history:
            # Job title and company
            doc.add_heading(f"{job.position} at {job.company}", level=2)
            
            # Dates and location
            date_text = f"{job.start_date} - {job.end_date}"
            if job.location:
                date_text += f" | {job.location}"
            doc.add_paragraph(date_text)
            
            # Description
            if job.description:
                for desc in job.description:
                    doc.add_paragraph(f"• {desc}", style='List Bullet')
            
            # Achievements
            if job.achievements:
                for achievement in job.achievements:
                    doc.add_paragraph(f"• {achievement}", style='List Bullet')
            
            doc.add_paragraph()  # Spacing
    
    def _add_docx_skills(self, doc: Document, cv_data: CVData):
        """Add skills to DOCX document"""
        doc.add_heading("SKILLS", level=1)
        
        if cv_data.technical_skills:
            doc.add_heading("Technical Skills:", level=2)
            tech_skills = []
            for skill in cv_data.technical_skills:
                skill_text = skill.name
                if skill.proficiency:
                    skill_text += f" ({skill.proficiency})"
                tech_skills.append(skill_text)
            doc.add_paragraph(', '.join(tech_skills))
        
        if cv_data.soft_skills:
            doc.add_heading("Soft Skills:", level=2)
            soft_skills = []
            for skill in cv_data.soft_skills:
                skill_text = skill.name
                if skill.proficiency:
                    skill_text += f" ({skill.proficiency})"
                soft_skills.append(skill_text)
            doc.add_paragraph(', '.join(soft_skills))
    
    def _add_docx_education(self, doc: Document, education: List):
        """Add education to DOCX document"""
        doc.add_heading("EDUCATION", level=1)
        
        for edu in education:
            doc.add_heading(f"{edu.degree} in {edu.field_of_study}", level=2)
            doc.add_paragraph(f"{edu.institution} | {edu.graduation_date}")
            if edu.gpa:
                doc.add_paragraph(f"GPA: {edu.gpa}")
            doc.add_paragraph()  # Spacing
    
    def _add_docx_projects(self, doc: Document, projects: List):
        """Add projects to DOCX document"""
        doc.add_heading("PROJECTS", level=1)
        
        for project in projects:
            doc.add_heading(project.name, level=2)
            if project.technologies:
                doc.add_paragraph(f"Technologies: {', '.join(project.technologies)}")
            if project.description:
                doc.add_paragraph(project.description)
            doc.add_paragraph()  # Spacing
    
    def _create_text_header(self, cv_data: CVData) -> str:
        """Create text header"""
        header_lines = []
        header_lines.append(cv_data.contact_info.full_name.upper())
        header_lines.append("=" * len(cv_data.contact_info.full_name))
        
        contact_info = []
        if cv_data.contact_info.email:
            contact_info.append(cv_data.contact_info.email)
        if cv_data.contact_info.phone:
            contact_info.append(cv_data.contact_info.phone)
        if cv_data.contact_info.location:
            contact_info.append(cv_data.contact_info.location)
        
        if contact_info:
            header_lines.append(' | '.join(contact_info))
        
        online_info = []
        if cv_data.contact_info.linkedin:
            online_info.append(f"LinkedIn: {cv_data.contact_info.linkedin}")
        if cv_data.contact_info.github:
            online_info.append(f"GitHub: {cv_data.contact_info.github}")
        if cv_data.contact_info.portfolio:
            online_info.append(f"Portfolio: {cv_data.contact_info.portfolio}")
        
        if online_info:
            header_lines.append(' | '.join(online_info))
        
        return '\n'.join(header_lines) 